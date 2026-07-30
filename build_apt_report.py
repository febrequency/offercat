from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches


OUT_DIR = Path("outputs/apt_unit_research")
OUT_FILE = OUT_DIR / "2026APT学员单位调研报告.docx"
FONT_NAME = "STSong"


SOURCES = {
    "mptc": "https://mptc.gov.kh/en/mptcs-overview/",
    "mptc_home": "https://mptc.gov.kh/en/",
    "cadt_about": "https://cadt.edu.kh/about/",
    "cadt_rebrand": "https://cadt.edu.kh/announcement-niptict-national-institute-rebranded-to-cambodia-academy-of-digital-technology-cadt/",
    "laotrade": "https://www.laotradeportal.gov.la/en-gb/site/display/2634",
    "kpl": "https://kpl.gov.la/EN/detail.aspx?id=96444",
    "malaysia_spab": "https://komunikasi.spab.gov.my/",
    "malaysia_resource": "https://pusatsumber.komunikasi.gov.my/visi-misi-dan-fungsi",
    "myanmar_portal": "https://myanmar.gov.mm/ministry-of-digital-development-and-communication",
    "myanmar_moi_links": "https://www.moi.gov.mm/moi%3Aeng/content/websites-ministries",
    "moitt": "https://moitt.gov.pk/Overview",
    "trc": "https://www.trc.gov.lk/complaint/pages_e.php?id=2",
    "slt_profile": "https://www.slt.lk/en/content/profile",
    "slt_corp": "https://training.slt.lk/reports-html/annual/2019/supplementary_information/corporate_information.html",
    "etda": "https://www.etda.or.th/en/about-etda/mission-of-etda.aspx",
    "nt_oneweb": "https://www.ntplc.co.th/en/news/update-news/nt-eutelsat-oneweb-snp-gateway-en",
    "nt_ctrls": "https://www.ntplc.co.th/en/news/update-news/news29",
    "mst_vtf": "https://mst.gov.vn/quy-dich-vu-vien-thong-cong-ich-viet-nam-197250905094109177.htm",
    "mst_vtf_2026": "https://mst.gov.vn/quy-dinh-moi-ve-to-chuc-hoat-dong-quy-dich-vu-vien-thong-cong-ich-viet-nam-197260130141357389.htm",
    "aiti": "https://www.aiti.gov.bn/about-us/who-we-are/",
}


ENTRIES = [
    {
        "cn": "柬埔寨邮电部",
        "en": "Ministry of Post and Telecommunications",
        "research_cn": (
            "柬埔寨邮电部是柬埔寨负责邮政、电信和信息通信技术相关事务的政府部门。"
            "其官网将愿景表述为促进有效网络基础设施互联和可及的邮政、电信、ICT服务，"
            "以服务柬埔寨国内、区域和国际连接，并支持社会经济发展与减贫。官网列明的任务包括："
            "管理邮政、电信和ICT领域，扩大可靠、安全的骨干基础设施，推动邮政、电信和ICT现代化，"
            "发展国家及地方层面的数字化公共服务，加强ICT认知、培训和能力建设，并推动公平竞争和消费者权益保护。"
            f"信息源：{SOURCES['mptc']}；{SOURCES['mptc_home']}"
        ),
        "research_en": (
            "The Ministry of Post and Telecommunications is Cambodia's government ministry responsible for the postal, "
            "telecommunications and ICT sectors. Its official website states a vision of promoting effective network "
            "infrastructure connectivity and accessible postal, telecommunications and ICT services for national, regional "
            "and global connectivity, contributing to socioeconomic development and poverty reduction. Its stated missions "
            "include sector management, expansion of reliable and secure backbone infrastructure, modernization of postal, "
            "telecommunications and ICT services, digital development of public services, ICT awareness and capacity building, "
            "and the protection of fair competition and consumer rights. "
            f"Sources: {SOURCES['mptc']}; {SOURCES['mptc_home']}"
        ),
    },
    {
        "cn": "柬埔寨数字技术学院",
        "en": "Cambodia Academy of Digital Technology",
        "research_cn": (
            "柬埔寨数字技术学院（CADT）是柬埔寨面向数字技术和创新的公共教育与研究机构。"
            "CADT官网介绍其定位为推动柬埔寨形成活力、包容和可持续数字社会的公共研究教育机构；"
            "其使命包括提供数字技术高等教育和专业培训、开展研究与知识推进、促进创新创业，"
            "服务数字政府、数字经济和数字社会建设。官网还显示CADT下设数字技术学院、数字治理学院、"
            "数字研究与创新学院。柬埔寨邮电部官网也将CADT列为其自治单位之一。"
            f"信息源：{SOURCES['cadt_about']}；{SOURCES['mptc_home']}"
        ),
        "research_en": (
            "Cambodia Academy of Digital Technology (CADT) is a public education and research institution focused on "
            "digital technology and innovation. CADT's official website describes it as a public research and education "
            "institution dedicated to accelerating a vibrant, inclusive and sustainable digital society in Cambodia. Its "
            "missions cover higher education and professional training in digital technology, research and advancement of "
            "knowledge, and innovation and entrepreneurship for better academic and social services. CADT comprises the "
            "Institute of Digital Technology, Institute of Digital Governance, and Institute of Digital Research and Innovation. "
            "The Ministry of Post and Telecommunications also lists CADT among its autonomous units. "
            f"Sources: {SOURCES['cadt_about']}; {SOURCES['mptc_home']}"
        ),
    },
    {
        "cn": "柬埔寨数字技术学院",
        "en": "Cambodia Academy of Digital Technology",
        "research_cn": (
            "该单位在名单中对应两名学员，故本报告按原名单保留第二个条目。"
            "柬埔寨数字技术学院（CADT）由原NIPTICT于2021年更名扩展而来，官方公告称其分设数字技术、"
            "数字治理、数字研究与创新三个下属学院，并以培养数字人才和创新者、推动柬埔寨迈向数字社会为目标。"
            "CADT官网当前介绍也强调其教育、研究、创新三项核心任务，服务数字政府、数字经济和数字社会。"
            f"信息源：{SOURCES['cadt_rebrand']}；{SOURCES['cadt_about']}"
        ),
        "research_en": (
            "This organization appears twice in the trainee list, so the report keeps a second entry for alignment with the source list. "
            "According to CADT's official announcement, the former NIPTICT was rebranded and expanded into CADT in 2021, with three "
            "subordinate institutes: Digital Technology, Digital Governance, and Digital Research and Innovation. The official CADT site "
            "also emphasizes its three core missions of education and professional training, research, and innovation, all supporting "
            "digital government, the digital economy and digital society in Cambodia. "
            f"Sources: {SOURCES['cadt_rebrand']}; {SOURCES['cadt_about']}"
        ),
    },
    {
        "cn": "老挝技术和通信部",
        "en": "Ministry of Technology and Communications",
        "research_cn": (
            "老挝技术和通信部是老挝负责技术、通信、数字化及相关监管工作的政府部门。"
            "老挝官方贸易门户收录的2022年部令显示，该部依据其组织和职能法令负责电信和ICT设备管理，"
            "包括检验、型号核准、合格声明和技术标准标签等，以保障设备质量、标准、安全和社会公共利益。"
            "老挝官方通讯社2026年报道显示，该部推动数字政府、数字经济和数字社会建设，并将数字基础设施、"
            "数字人才培养、人工智能应用和公共服务数字化作为重点方向。"
            f"信息源：{SOURCES['laotrade']}；{SOURCES['kpl']}"
        ),
        "research_en": (
            "The Ministry of Technology and Communications is the Lao government ministry responsible for technology, communications, "
            "digital development and related regulatory work. A 2022 ministerial decision published on the official Lao Trade Portal "
            "shows that the ministry manages telecommunications and ICT equipment through inspection, type approval, conformity "
            "notification and technical standards labeling, with the purpose of ensuring quality, standards, security and public interest. "
            "A 2026 report by the Lao official news agency also describes the ministry's work on digital government, digital economy "
            "and digital society, including priorities in digital infrastructure, digital workforce development, AI application and "
            "digital public services. "
            f"Sources: {SOURCES['laotrade']}; {SOURCES['kpl']}"
        ),
    },
    {
        "cn": "马来西亚通信部",
        "en": "Ministry of Communications",
        "research_cn": (
            "马来西亚通信部是马来西亚联邦政府通信相关事务的主管部门之一。其官方投诉/咨询系统说明，"
            "该部受理范围包括广播业、涵盖电话和互联网服务的电信业，以及邮政和快递服务。"
            "该部资源中心官网还将通信、多媒体、广播、电影、本地音乐产业及相关领域列为核心信息领域。"
            "据此可见，该部职责重点围绕通信与媒体相关公共服务、行业事务和信息服务。"
            f"信息源：{SOURCES['malaysia_spab']}；{SOURCES['malaysia_resource']}"
        ),
        "research_en": (
            "Malaysia's Ministry of Communications is a federal government ministry involved in communications-related affairs. "
            "Its official complaint and inquiry system states that the ministry's complaint scope covers the broadcasting industry, "
            "telecommunications including telephone and internet services, and postal and courier services. The ministry resource "
            "centre also identifies communications, multimedia, broadcasting, film, the local music industry and related fields as "
            "core information areas. On that basis, the ministry's work is centered on communications and media-related public services, "
            "industry matters and information services. "
            f"Sources: {SOURCES['malaysia_spab']}; {SOURCES['malaysia_resource']}"
        ),
    },
    {
        "cn": "缅甸数字发展与通信部",
        "en": "Ministry of Digital Development and Communications",
        "research_cn": (
            "缅甸国家门户列有“数字发展与通信部”页面，缅甸信息部英文页面也将该部列入政府部委网站清单。"
            "该官方清单显示，该部关联机构包括邮政与电信部门、信息技术与网络安全部门、缅甸邮政电信和缅甸邮政等。"
            "因此，从官方门户可确认该部覆盖数字发展、通信、邮政电信以及信息技术和网络安全相关职能。"
            f"信息源：{SOURCES['myanmar_portal']}；{SOURCES['myanmar_moi_links']}"
        ),
        "research_en": (
            "Myanmar's national portal contains a page for the Ministry of Digital Development and Communications, and the English "
            "website of Myanmar's Ministry of Information lists it among government ministry websites. The official list links the ministry "
            "with the Posts and Telecommunications Department, the Department of Information Technology and Cyber Security, Myanma Posts "
            "and Telecommunications, and Myanmar Post. This confirms, from official portal information, that the ministry covers digital "
            "development, communications, postal and telecommunications affairs, and information technology and cyber security functions. "
            f"Sources: {SOURCES['myanmar_portal']}; {SOURCES['myanmar_moi_links']}"
        ),
    },
    {
        "cn": "巴基斯坦信息技术和电信部",
        "en": "Ministry of IT & Telecom",
        "research_cn": (
            "巴基斯坦信息技术和电信部（MoITT）是巴基斯坦政府负责信息技术和电信事务的国家牵头部门。"
            "其官网称，该部是政府在规划、协调和指导IT与电信项目方面的国家焦点部门和赋能机构，"
            "目标是通过ICT应用和电信平台支持知识型经济和公共服务改善。其授权事项包括制定IT和电信总体规划与政策、"
            "协调各省政府和相关机构、促进IT应用和人力资源发展、制定电信政策和立法，并处理与PTA、FAB等机构相关的联邦政府职能。"
            f"信息源：{SOURCES['moitt']}"
        ),
        "research_en": (
            "Pakistan's Ministry of IT & Telecom (MoITT) is the national focal ministry for information technology and telecommunications. "
            "Its official website describes the ministry as the Government of Pakistan's national focal ministry and enabling arm for planning, "
            "coordinating and directing IT and telecommunications programs and projects for economic development. Its vision and mission focus "
            "on ICT applications, telecom platforms, a knowledge-based economy, policy and legal frameworks, ICT infrastructure and improved "
            "public services. Its mandate includes IT and telecom planning and policy, coordination with provincial governments and other bodies, "
            "IT applications and human-resource development, telecom policy and legislation, and federal functions related to PTA and FAB. "
            f"Source: {SOURCES['moitt']}"
        ),
    },
    {
        "cn": "斯里兰卡电信监管委员会",
        "en": "Telecommunications Regulatory Commission of Sri Lanka",
        "research_cn": (
            "斯里兰卡电信监管委员会（TRCSL）是斯里兰卡国家电信监管机构，依据1996年《斯里兰卡电信法修正案》设立。"
            "其官网说明，TRCSL目标包括确保斯里兰卡国内和国际电信服务可靠高效，保护消费者、购买者、用户和公众利益，"
            "并维护有效竞争。其职能涵盖电信系统、无线电频率、私人网络、供应商及布线等许可处理，资费监管，合规监管，"
            "频谱利用监测，消费者投诉处理，以及向公众提供服务质量与服务种类信息。"
            f"信息源：{SOURCES['trc']}"
        ),
        "research_en": (
            "The Telecommunications Regulatory Commission of Sri Lanka (TRCSL) is Sri Lanka's national telecommunications regulatory agency, "
            "established under the Sri Lanka Telecommunication (Amendment) Act No. 27 of 1996. Its official website states objectives including "
            "reliable and efficient national and international telecommunications services, protection of consumers, purchasers, users and the "
            "public interest, and maintenance of effective competition. Its functions cover licensing for telecommunication systems, radio "
            "frequencies, private networks, vendors and cabling works, tariff regulation, compliance monitoring, spectrum management, consumer "
            "complaints and public information on service quality and variety. "
            f"Source: {SOURCES['trc']}"
        ),
    },
    {
        "cn": "斯里兰卡电信公共有限公司",
        "en": "Sri Lanka Telecom PLC",
        "research_cn": (
            "斯里兰卡电信公共有限公司（Sri Lanka Telecom PLC，SLT）是斯里兰卡国家ICT解决方案提供商，"
            "也是该国主要宽带和骨干基础设施服务提供商。公司官网介绍，SLT服务斯里兰卡连接需求超过163年，"
            "通过光纤、铜缆和无线接入网络服务超过900万客户，并提供固定和移动电话、宽带、数据服务、IPTV、"
            "云计算、托管和网络解决方案等多元ICT服务。其官方年报信息显示，公司于1996年改制为公众有限责任公司，"
            "并于2003年在科伦坡证券交易所上市。"
            f"信息源：{SOURCES['slt_profile']}；{SOURCES['slt_corp']}"
        ),
        "research_en": (
            "Sri Lanka Telecom PLC (SLT) is Sri Lanka's national ICT solutions provider and a leading broadband and backbone infrastructure "
            "services provider. SLT's official profile says it has served the country's connectivity needs for more than 163 years, supports "
            "more than nine million customers through fiber, copper and wireless access networks, and provides diversified ICT services including "
            "fixed and mobile telephony, broadband, data services, IPTV, cloud computing, hosting and networking solutions. Official corporate "
            "information also states that SLT became a public limited liability company in 1996 and was listed on the Colombo Stock Exchange in 2003. "
            f"Sources: {SOURCES['slt_profile']}; {SOURCES['slt_corp']}"
        ),
    },
    {
        "cn": "泰国电子交易发展局（ETDA）",
        "en": "ETDA Electronic Transactions Development Agency",
        "research_cn": (
            "泰国电子交易发展局（ETDA）成立于2011年，是泰国推动电子交易和数字治理的重要机构。"
            "其官网称，ETDA依据《电子交易法》和《电子交易发展局法》承担促进、支持和发展电子交易、"
            "线上交易及相关数字服务的职责，目标是让政府、企业和公众能够以可靠、安全的方式开展线上交易。"
            "ETDA工作机制包括许可、注册、通知、标准制定、立法和沙盒测试，并开展数据分析、人才培养、咨询、"
            "防欺诈和创新促进等基础工作。"
            f"信息源：{SOURCES['etda']}"
        ),
        "research_en": (
            "Thailand's Electronic Transactions Development Agency (ETDA) was founded in 2011 and is a key agency for electronic transactions "
            "and digital governance. Its official website states that ETDA promotes, supports and develops electronic transactions, online "
            "transactions and related digital services under the Electronic Transactions Act and the Electronic Transactions Development Agency Act. "
            "Its purpose is to enable government, businesses and the public to conduct trusted, secure and safe online transactions. ETDA's working "
            "mechanisms include licensing, registration, notification, standard-setting, legislation and sandbox testing, alongside data analysis, "
            "personnel development, consultation, fraud prevention and innovation promotion. "
            f"Source: {SOURCES['etda']}"
        ),
    },
    {
        "cn": "泰国国家电信公共有限公司（NT）",
        "en": "National Telecom Public Company Limited (NT)",
        "research_cn": (
            "泰国国家电信公共有限公司（NT）是泰国国有电信企业，隶属于数字经济与社会部（MDES）。"
            "NT官网新闻介绍其为电信和卫星通信服务主要提供者，致力于推进数字基础设施，服务政府机构、企业和公众。"
            "NT官方信息还显示，公司于2021年由CAT Telecom与TOT Public Company Limited合并并公司化成立，"
            "业务涵盖硬基础设施、国际电信基础设施、国际互联网网关、固定宽带、移动网络服务以及数字基础设施与服务。"
            f"信息源：{SOURCES['nt_oneweb']}；{SOURCES['nt_ctrls']}"
        ),
        "research_en": (
            "National Telecom Public Company Limited (NT) is a Thai state-owned telecommunications enterprise under the Ministry of Digital Economy "
            "and Society (MDES). NT's official news describes it as a leading provider of telecommunications and satellite communication services "
            "committed to advancing digital infrastructure for government agencies, businesses and the public. Official NT information also states "
            "that the company was established and corporatized in 2021 from the merger of CAT Telecom and TOT Public Company Limited, with business "
            "covering hard infrastructure, international telecommunications infrastructure, international internet gateways, fixed broadband, mobile "
            "network services, and digital infrastructure and services. "
            f"Sources: {SOURCES['nt_oneweb']}; {SOURCES['nt_ctrls']}"
        ),
    },
    {
        "cn": "越南公用电信服务基金（隶属越南科学技术部）",
        "en": "Vietnam Public-utility Telecommunication Service Fund, Ministry of Science and Technology",
        "research_cn": (
            "越南公用电信服务基金（VTF）是越南科学技术部下属国家财政性质机构。"
            "越南科学技术部官网介绍，该基金现隶属于科学技术部，并列明办公地址、联系方式和网站。"
            "该部2026年发布的信息显示，根据第06/2026/QĐ-TTg号决定，VTF为由科学技术部管理的预算外国家财政基金，"
            "按公立事业单位模式运作，主要用于落实国家公共电信政策，承担资金管理、支持公共电信服务、"
            "组织电信企业缴费、接收合法资金来源、监督资金使用和履行财务会计公开报告等任务。"
            f"信息源：{SOURCES['mst_vtf']}；{SOURCES['mst_vtf_2026']}"
        ),
        "research_en": (
            "The Vietnam Public-utility Telecommunication Service Fund (VTF) is a state financial institution under Vietnam's Ministry of Science "
            "and Technology. The ministry's official website identifies the fund as directly under the ministry and provides its office address, "
            "contact details and website. A 2026 ministry notice on Decision No. 06/2026/QD-TTg states that VTF is a non-budgetary state financial "
            "fund managed by the Ministry of Science and Technology and operating as a public service unit. It supports the implementation of state "
            "policy on public telecommunications, manages funds for public-utility telecommunications, collects financial contributions from telecom "
            "enterprises, receives lawful financial sources, monitors use of supported funds, and implements financial, accounting, transparency and "
            "reporting obligations. "
            f"Sources: {SOURCES['mst_vtf']}; {SOURCES['mst_vtf_2026']}"
        ),
    },
    {
        "cn": "文莱信息通信技术产业管理局（AITI）",
        "en": "Authority for Infocommunications Technology Industry of Brunei Darussalam (AITI)",
        "research_cn": (
            "文莱信息通信技术产业管理局（AITI）是依据2001年AITI令设立、于2003年1月1日成立的法定机构。"
            "其官网列明主要职责包括监管电信和邮政系统及服务，发展本地ICT产业，通过促进有竞争力的价格和服务质量保护消费者利益，"
            "管理国家无线电频谱，并执行2025年个人数据保护令。AITI还作为政府政策和监管框架的咨询机构、"
            "行业协作和创新的促进者，以及推动数字产业发展项目的实施者，服务“数字文莱”建设。"
            f"信息源：{SOURCES['aiti']}"
        ),
        "research_en": (
            "The Authority for Infocommunications Technology Industry of Brunei Darussalam (AITI) is a statutory body established on "
            "1 January 2003 under the AITI Order, 2001. Its official website lists key responsibilities including regulatory functions for "
            "telecommunications and postal systems and services, development of the local ICT industry, protection of consumer interests through "
            "competitive pricing and service quality, management of the national radio-frequency spectrum, and administration and enforcement of "
            "the Personal Data Protection Order, 2025. AITI also acts as an adviser to government on policy and regulatory frameworks, a facilitator "
            "for collaboration and innovation, and a champion of development programmes supporting Digital Brunei. "
            f"Source: {SOURCES['aiti']}"
        ),
    },
]


def set_run_font(run, bold=None):
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    if bold is not None:
        run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), FONT_NAME)


def apply_expanded_content():
    expanded = [
        {
            "research_cn": (
                "柬埔寨邮电部是柬埔寨负责邮政、电信和信息通信技术相关事务的中央政府部门，也是该国推进数字连接和通信基础设施建设的重要主管机关。"
                "其官网将愿景表述为促进有效网络基础设施互联和可及的邮政、电信、ICT服务，以服务柬埔寨国内、区域和国际连接，并支持社会经济发展与减贫。"
                "官网列明的任务包括管理邮政、电信和ICT领域，扩大可靠、安全的骨干基础设施，推动邮政、电信和ICT现代化，发展国家及地方层面的数字化公共服务，"
                "加强ICT认知、培训和能力建设，并推动公平竞争和消费者权益保护。该部官网还显示其下设和关联电信监管、邮政、数字技术教育培训等单位，说明其职责不仅限于行业行政管理，"
                "也覆盖数字人才、公共服务数字化和国家通信能力建设等方面。信息源："
                f"{SOURCES['mptc']}；{SOURCES['mptc_home']}"
            ),
            "research_en": (
                "The Ministry of Post and Telecommunications is Cambodia's central government ministry responsible for postal, telecommunications "
                "and ICT affairs, and it is a key authority for digital connectivity and communications infrastructure. Its official website states "
                "a vision of promoting effective network infrastructure connectivity and accessible postal, telecommunications and ICT services for "
                "national, regional and global connectivity, while contributing to socioeconomic development and poverty reduction. Its stated missions "
                "include sector management, expansion of reliable and secure backbone infrastructure, modernization of postal, telecommunications and "
                "ICT services, digital development of public services at national and local levels, ICT awareness and capacity building, fair competition "
                "and consumer protection. The ministry's website also shows links with regulatory, postal and digital-technology education institutions, "
                "which indicates a remit extending from administration to digital talent, public-service digitization and national communications capacity. "
                f"Sources: {SOURCES['mptc']}; {SOURCES['mptc_home']}"
            ),
        },
        {
            "research_cn": (
                "柬埔寨数字技术学院（CADT）是柬埔寨面向数字技术和创新的公共教育与研究机构，在柬埔寨数字人才培养体系中具有基础性作用。"
                "CADT官网介绍其定位为推动柬埔寨形成活力、包容和可持续数字社会的公共研究教育机构；其使命包括提供数字技术高等教育和专业培训、开展研究与知识推进、"
                "促进创新创业，服务数字政府、数字经济和数字社会建设。官网显示，CADT下设数字技术学院、数字治理学院、数字研究与创新学院，覆盖技术教育、公共部门数字治理、"
                "应用研究和创新孵化等方向。柬埔寨邮电部官网也将CADT列为其自治单位之一，说明CADT与国家数字政策和通信技术人才建设之间联系紧密，既承担教学培训功能，"
                "也参与支撑政府数字化转型和产业创新生态建设。信息源："
                f"{SOURCES['cadt_about']}；{SOURCES['mptc_home']}"
            ),
            "research_en": (
                "Cambodia Academy of Digital Technology (CADT) is a public education and research institution focused on digital technology and "
                "innovation, playing a foundational role in Cambodia's digital talent pipeline. CADT's official website describes it as a public "
                "research and education institution dedicated to accelerating a vibrant, inclusive and sustainable digital society in Cambodia. Its "
                "missions cover higher education and professional training in digital technology, research and advancement of knowledge, and innovation "
                "and entrepreneurship for digital government, the digital economy and digital society. CADT comprises the Institute of Digital Technology, "
                "Institute of Digital Governance, and Institute of Digital Research and Innovation, covering technical education, public-sector digital "
                "governance, applied research and innovation incubation. The Ministry of Post and Telecommunications also lists CADT among its autonomous "
                "units, showing its close connection with national digital policy, communications technology capacity building and digital transformation. "
                f"Sources: {SOURCES['cadt_about']}; {SOURCES['mptc_home']}"
            ),
        },
        {
            "research_cn": (
                "柬埔寨数字技术学院（CADT）由原NIPTICT于2021年更名扩展而来，官方公告称其分设数字技术、数字治理、数字研究与创新三个下属学院，"
                "并以培养数字人才和创新者、推动柬埔寨迈向数字社会为目标。CADT官网当前介绍也强调其教育、研究、创新三项核心任务，服务数字政府、数字经济和数字社会。"
                "从机构设置看，数字技术学院侧重ICT和数字技术人才培养，数字治理学院侧重公共部门数字化和治理能力建设，数字研究与创新学院则面向研究、创新和应用转化。"
                "这使CADT同时具备高校、培训机构和政策支撑平台的特征，能够为政府部门、通信产业和数字经济发展提供人才、研究和创新服务。信息源："
                f"{SOURCES['cadt_rebrand']}；{SOURCES['cadt_about']}"
            ),
            "research_en": (
                "According to CADT's official announcement, the former NIPTICT was rebranded and expanded into Cambodia Academy of Digital Technology "
                "in 2021, with three subordinate institutes: Digital Technology, Digital Governance, and Digital Research and Innovation. The official "
                "CADT site also emphasizes its three core missions of education and professional training, research, and innovation, all supporting "
                "digital government, the digital economy and digital society in Cambodia. In institutional terms, the Institute of Digital Technology "
                "focuses on ICT and digital skills, the Institute of Digital Governance supports public-sector digitization and governance capacity, "
                "and the Institute of Digital Research and Innovation focuses on research, innovation and application. This gives CADT the combined "
                "character of a higher-education institution, training body and policy-support platform for government, the communications sector and "
                "the broader digital economy. "
                f"Sources: {SOURCES['cadt_rebrand']}; {SOURCES['cadt_about']}"
            ),
        },
        {
            "research_cn": (
                "老挝技术和通信部是老挝负责技术、通信、数字化及相关监管工作的政府部门，职责范围与电信基础设施、ICT设备管理、数字政府和数字经济建设密切相关。"
                "老挝官方贸易门户收录的2022年部令显示，该部依据其组织和职能法令负责电信和ICT设备管理，包括检验、型号核准、合格声明和技术标准标签等，"
                "以保障设备质量、标准、安全和社会公共利益。这说明该部承担通信设备准入、技术标准和市场监管职能。老挝官方通讯社2026年报道显示，该部推动数字政府、"
                "数字经济和数字社会建设，并将数字基础设施、数字人才培养、人工智能应用和公共服务数字化作为重点方向。整体看，该部既是通信技术行政主管部门，"
                "也是老挝数字转型战略落地的重要执行机构。信息源："
                f"{SOURCES['laotrade']}；{SOURCES['kpl']}"
            ),
            "research_en": (
                "The Ministry of Technology and Communications is the Lao government ministry responsible for technology, communications, digital "
                "development and related regulation, with a remit closely linked to telecommunications infrastructure, ICT equipment management, "
                "digital government and the digital economy. A 2022 ministerial decision published on the official Lao Trade Portal shows that the "
                "ministry manages telecommunications and ICT equipment through inspection, type approval, conformity notification and technical "
                "standards labeling, with the purpose of ensuring quality, standards, security and public interest. This indicates responsibilities "
                "in equipment access, technical standards and market regulation. A 2026 report by the Lao official news agency also describes the "
                "ministry's work on digital government, digital economy and digital society, including priorities in digital infrastructure, digital "
                "workforce development, AI application and digital public services. Overall, the ministry is both a communications-technology regulator "
                "and an implementation body for Laos' digital transformation agenda. "
                f"Sources: {SOURCES['laotrade']}; {SOURCES['kpl']}"
            ),
        },
        {
            "research_cn": (
                "马来西亚通信部是马来西亚联邦政府通信相关事务的主管部门之一，工作范围覆盖通信服务、媒体传播、广播、邮政快递和公众信息服务等领域。"
                "其官方投诉/咨询系统说明，该部受理范围包括广播业、涵盖电话和互联网服务的电信业，以及邮政和快递服务，体现其在通信服务质量、行业协调和公众诉求处理方面的职责。"
                "该部资源中心官网还将通信、多媒体、广播、电影、本地音乐产业及相关领域列为核心信息领域，说明其不仅处理电信行业事务，也涉及内容传播、媒体产业和公共沟通。"
                "从官方信息看，该部在马来西亚数字社会建设中的角色主要是协调通信基础服务、媒体传播秩序和面向公众的信息服务，支撑政府与社会之间的信息流通和数字连接能力。信息源："
                f"{SOURCES['malaysia_spab']}；{SOURCES['malaysia_resource']}"
            ),
            "research_en": (
                "Malaysia's Ministry of Communications is one of the federal government ministries responsible for communications-related affairs, "
                "covering communications services, media, broadcasting, postal and courier services, and public information functions. Its official "
                "complaint and inquiry system states that the ministry's complaint scope covers the broadcasting industry, telecommunications including "
                "telephone and internet services, and postal and courier services, reflecting its role in service quality, sector coordination and public "
                "complaint handling. The ministry resource centre also identifies communications, multimedia, broadcasting, film, the local music industry "
                "and related fields as core information areas, showing that the ministry deals not only with telecommunications but also with content, "
                "media industries and public communication. On the basis of official information, its role in Malaysia's digital society is to coordinate "
                "basic communications services, media order and public information services, supporting information flows and digital connectivity. "
                f"Sources: {SOURCES['malaysia_spab']}; {SOURCES['malaysia_resource']}"
            ),
        },
        {
            "research_cn": (
                "缅甸数字发展与通信部是缅甸政府体系中负责数字发展和通信相关事务的部门。缅甸国家门户列有“数字发展与通信部”页面，缅甸信息部英文页面也将该部列入政府部委网站清单。"
                "该官方清单显示，该部关联机构包括邮政与电信部门、信息技术与网络安全部门、缅甸邮政电信和缅甸邮政等。由此可见，该部的工作范围覆盖通信网络、邮政服务、"
                "信息技术管理和网络安全等多个方面，既包括传统电信与邮政基础服务，也包括数字发展和网络安全能力建设。由于相关官方公开介绍较为简略，本报告以政府门户和信息部官网清单为依据，"
                "保守确认其机构定位：该部是缅甸推进数字发展、通信服务、邮政电信管理和信息技术安全相关工作的政府主管部门。信息源："
                f"{SOURCES['myanmar_portal']}；{SOURCES['myanmar_moi_links']}"
            ),
            "research_en": (
                "Myanmar's Ministry of Digital Development and Communications is the government ministry responsible for digital development and "
                "communications-related affairs. Myanmar's national portal contains a page for the ministry, and the English website of Myanmar's "
                "Ministry of Information lists it among government ministry websites. The official list links the ministry with the Posts and "
                "Telecommunications Department, the Department of Information Technology and Cyber Security, Myanma Posts and Telecommunications, "
                "and Myanmar Post. This indicates a scope covering communications networks, postal services, information technology management and "
                "cyber security, combining traditional telecommunications and postal services with digital development and security capacity building. "
                "Because the available official public description is brief, this report conservatively identifies the ministry, based on government "
                "portal and Ministry of Information sources, as the authority for digital development, communications services, postal and telecom "
                "management, and information-technology security functions in Myanmar. "
                f"Sources: {SOURCES['myanmar_portal']}; {SOURCES['myanmar_moi_links']}"
            ),
        },
        {
            "research_cn": (
                "巴基斯坦信息技术和电信部（MoITT）是巴基斯坦政府负责信息技术和电信事务的国家牵头部门。其官网称，该部是政府在规划、协调和指导IT与电信项目方面的国家焦点部门和赋能机构，"
                "目标是通过ICT应用和电信平台支持知识型经济和公共服务改善。其愿景和使命强调通过ICT基础设施、政策法律框架、电子政务和公共服务数字化推动经济社会发展。"
                "官网列明的授权事项包括制定IT和电信总体规划与政策，协调各省政府和相关机构，促进IT应用和人力资源发展，制定电信政策和立法，并处理与PTA、FAB等机构相关的联邦政府职能。"
                "因此，MoITT既承担数字政策制定、行业协调和监管框架建设，也负责推动ICT产业发展、数字技能和政府数字化应用。信息源："
                f"{SOURCES['moitt']}"
            ),
            "research_en": (
                "Pakistan's Ministry of IT & Telecom (MoITT) is the national focal ministry for information technology and telecommunications. Its "
                "official website describes the ministry as the Government of Pakistan's national focal ministry and enabling arm for planning, "
                "coordinating and directing IT and telecommunications programs and projects for economic development. Its vision and mission emphasize "
                "ICT infrastructure, policy and legal frameworks, e-government and digitized public services as tools for socioeconomic development. "
                "Its mandate includes IT and telecom planning and policy, coordination with provincial governments and other bodies, promotion of IT "
                "applications and human-resource development, telecom policy and legislation, and federal functions related to PTA and FAB. Therefore, "
                "MoITT combines digital policy-making, sector coordination and regulatory-framework development with ICT industry promotion, digital "
                "skills and government digital applications. "
                f"Source: {SOURCES['moitt']}"
            ),
        },
        {
            "research_cn": (
                "斯里兰卡电信监管委员会（TRCSL）是斯里兰卡国家电信监管机构，依据1996年《斯里兰卡电信法修正案》设立。其官网说明，TRCSL目标包括确保斯里兰卡国内和国际电信服务可靠高效，"
                "保护消费者、购买者、用户和公众利益，并维护有效竞争。官网列明的职能涵盖电信系统、无线电频率、私人网络、供应商及布线等许可处理，资费监管，合规监管，频谱利用监测，"
                "消费者投诉处理，以及向公众提供服务质量与服务种类信息。该机构的定位兼具行业准入、市场监管、频谱管理和消费者保护功能，是斯里兰卡电信市场秩序和通信资源配置的核心监管主体。"
                "对于涉及频谱、网络接入、服务质量和用户权益的问题，TRCSL均具有直接监管意义。信息源："
                f"{SOURCES['trc']}"
            ),
            "research_en": (
                "The Telecommunications Regulatory Commission of Sri Lanka (TRCSL) is Sri Lanka's national telecommunications regulatory agency, "
                "established under the Sri Lanka Telecommunication (Amendment) Act No. 27 of 1996. Its official website states objectives including "
                "reliable and efficient national and international telecommunications services, protection of consumers, purchasers, users and the "
                "public interest, and maintenance of effective competition. Its listed functions cover licensing for telecommunication systems, radio "
                "frequencies, private networks, vendors and cabling works, tariff regulation, compliance monitoring, spectrum-use monitoring, consumer "
                "complaints and public information on service quality and variety. TRCSL therefore combines market entry control, sector regulation, "
                "spectrum management and consumer protection. It is the central regulatory body for telecommunications market order and communications "
                "resource allocation in Sri Lanka, especially in matters involving spectrum, network access, service quality and user rights. "
                f"Source: {SOURCES['trc']}"
            ),
        },
        {
            "research_cn": (
                "斯里兰卡电信公共有限公司（Sri Lanka Telecom PLC，SLT）是斯里兰卡国家ICT解决方案提供商，也是该国主要宽带和骨干基础设施服务提供商。公司官网介绍，"
                "SLT服务斯里兰卡连接需求超过163年，通过光纤、铜缆和无线接入网络服务超过900万客户，并提供固定和移动电话、宽带、数据服务、IPTV、云计算、托管和网络解决方案等多元ICT服务。"
                "官网还强调SLT持续投资下一代宽带、企业ICT、数据中心和数字服务能力，面向个人用户、企业和公共部门提供连接与数字化解决方案。其官方年报信息显示，公司于1996年改制为公众有限责任公司，"
                "并于2003年在科伦坡证券交易所上市。整体看，SLT在斯里兰卡承担基础通信网络、宽带接入、企业数字服务和国家级ICT能力支撑等多重角色。信息源："
                f"{SOURCES['slt_profile']}；{SOURCES['slt_corp']}"
            ),
            "research_en": (
                "Sri Lanka Telecom PLC (SLT) is Sri Lanka's national ICT solutions provider and a leading broadband and backbone infrastructure "
                "services provider. SLT's official profile says it has served the country's connectivity needs for more than 163 years, supports more "
                "than nine million customers through fiber, copper and wireless access networks, and provides diversified ICT services including fixed "
                "and mobile telephony, broadband, data services, IPTV, cloud computing, hosting and networking solutions. The company also emphasizes "
                "investment in next-generation broadband, enterprise ICT, data centres and digital services for individual users, businesses and the "
                "public sector. Official corporate information states that SLT became a public limited liability company in 1996 and was listed on the "
                "Colombo Stock Exchange in 2003. Overall, SLT plays multiple roles in Sri Lanka: basic communications networks, broadband access, "
                "enterprise digital services and national ICT capability support. "
                f"Sources: {SOURCES['slt_profile']}; {SOURCES['slt_corp']}"
            ),
        },
        {
            "research_cn": (
                "泰国电子交易发展局（ETDA）成立于2011年，是泰国推动电子交易和数字治理的重要机构。其官网称，ETDA依据《电子交易法》和《电子交易发展局法》承担促进、支持和发展电子交易、"
                "线上交易及相关数字服务的职责，目标是让政府、企业和公众能够以可靠、安全的方式开展线上交易。ETDA工作机制包括许可、注册、通知、标准制定、立法和沙盒测试，并开展数据分析、"
                "人才培养、咨询、防欺诈和创新促进等基础工作。该机构关注的重点不是传统通信网络运营，而是数字交易生态的规则、信任、安全和服务创新，包括电子签名、数字身份、在线服务监管和消费者信任等方向。"
                "因此，ETDA在泰国数字经济中承担制度建设、标准规则、市场促进和数字信任基础设施支撑作用。信息源："
                f"{SOURCES['etda']}"
            ),
            "research_en": (
                "Thailand's Electronic Transactions Development Agency (ETDA) was founded in 2011 and is a key agency for electronic transactions "
                "and digital governance. Its official website states that ETDA promotes, supports and develops electronic transactions, online "
                "transactions and related digital services under the Electronic Transactions Act and the Electronic Transactions Development Agency Act. "
                "Its purpose is to enable government, businesses and the public to conduct trusted, secure and safe online transactions. ETDA's working "
                "mechanisms include licensing, registration, notification, standard-setting, legislation and sandbox testing, alongside data analysis, "
                "personnel development, consultation, fraud prevention and innovation promotion. Its focus is not traditional network operation but the "
                "rules, trust, safety and service innovation of the digital-transaction ecosystem, including electronic signatures, digital identity, "
                "online-service governance and consumer trust. ETDA therefore supports institutional development, standards, market promotion and digital "
                "trust infrastructure in Thailand's digital economy. "
                f"Source: {SOURCES['etda']}"
            ),
        },
        {
            "research_cn": (
                "泰国国家电信公共有限公司（NT）是泰国国有电信企业，隶属于数字经济与社会部（MDES）。NT官网新闻介绍其为电信和卫星通信服务主要提供者，致力于推进数字基础设施，"
                "服务政府机构、企业和公众。NT官方信息还显示，公司于2021年由CAT Telecom与TOT Public Company Limited合并并公司化成立，业务涵盖硬基础设施、国际电信基础设施、"
                "国际互联网网关、固定宽带、移动网络服务以及数字基础设施与服务。作为国有电信企业，NT既提供商业通信服务，也承担国家通信基础设施、国际连接能力和公共部门数字服务支撑任务。"
                "其卫星通信、宽带、国际网关和数字服务布局，说明NT在泰国数字经济基础设施和政府数字化应用中具有平台型作用。信息源："
                f"{SOURCES['nt_oneweb']}；{SOURCES['nt_ctrls']}"
            ),
            "research_en": (
                "National Telecom Public Company Limited (NT) is a Thai state-owned telecommunications enterprise under the Ministry of Digital Economy "
                "and Society (MDES). NT's official news describes it as a leading provider of telecommunications and satellite communication services "
                "committed to advancing digital infrastructure for government agencies, businesses and the public. Official NT information states that "
                "the company was established and corporatized in 2021 from the merger of CAT Telecom and TOT Public Company Limited, with business covering "
                "hard infrastructure, international telecommunications infrastructure, international internet gateways, fixed broadband, mobile network "
                "services, and digital infrastructure and services. As a state-owned telecom enterprise, NT provides commercial communications services while "
                "also supporting national communications infrastructure, international connectivity and public-sector digital services. Its satellite, "
                "broadband, international gateway and digital-service capabilities give it a platform role in Thailand's digital-economy infrastructure and "
                "government digitization. "
                f"Sources: {SOURCES['nt_oneweb']}; {SOURCES['nt_ctrls']}"
            ),
        },
        {
            "research_cn": (
                "越南公用电信服务基金（VTF）是越南科学技术部下属国家财政性质机构。越南科学技术部官网介绍，该基金现隶属于科学技术部，并列明办公地址、联系方式和网站。"
                "该部2026年发布的信息显示，根据第06/2026/QĐ-TTg号决定，VTF为由科学技术部管理的预算外国家财政基金，按公立事业单位模式运作，主要用于落实国家公共电信政策。"
                "其任务包括管理公共电信服务资金，支持公益电信服务项目，组织电信企业缴费，接收合法资金来源，监督资金使用，并履行财务会计、公开透明和报告义务。"
                "该基金的核心意义在于通过财政机制弥补商业电信服务不足，支持偏远地区、弱势群体或公共目标相关的通信服务供给，保障数字连接的普惠性和公共服务属性。信息源："
                f"{SOURCES['mst_vtf']}；{SOURCES['mst_vtf_2026']}"
            ),
            "research_en": (
                "The Vietnam Public-utility Telecommunication Service Fund (VTF) is a state financial institution under Vietnam's Ministry of Science "
                "and Technology. The ministry's official website identifies the fund as directly under the ministry and provides its office address, "
                "contact details and website. A 2026 ministry notice on Decision No. 06/2026/QD-TTg states that VTF is a non-budgetary state financial "
                "fund managed by the Ministry of Science and Technology and operating as a public service unit, mainly to implement state policy on public "
                "telecommunications. Its tasks include managing public-utility telecommunications funds, supporting public telecom service projects, "
                "collecting financial contributions from telecommunications enterprises, receiving lawful funding sources, monitoring the use of supported "
                "funds, and fulfilling financial, accounting, transparency and reporting obligations. The fund's core significance is to use financial "
                "mechanisms to address gaps left by purely commercial telecom services, supporting connectivity for remote areas, vulnerable groups and "
                "public-interest objectives. "
                f"Sources: {SOURCES['mst_vtf']}; {SOURCES['mst_vtf_2026']}"
            ),
        },
        {
            "research_cn": (
                "文莱信息通信技术产业管理局（AITI）是依据2001年AITI令设立、于2003年1月1日成立的法定机构，是文莱信息通信和邮政领域的重要监管与发展机构。"
                "其官网列明主要职责包括监管电信和邮政系统及服务，发展本地ICT产业，通过促进有竞争力的价格和服务质量保护消费者利益，管理国家无线电频谱，并执行2025年个人数据保护令。"
                "AITI还作为政府政策和监管框架的咨询机构、行业协作和创新的促进者，以及推动数字产业发展项目的实施者，服务“数字文莱”建设。该机构兼具监管、产业发展、消费者保护、频谱资源管理和数据保护职责，"
                "在文莱数字经济、通信市场治理和个人数据保护制度建设中具有综合平台作用。信息源："
                f"{SOURCES['aiti']}"
            ),
            "research_en": (
                "The Authority for Infocommunications Technology Industry of Brunei Darussalam (AITI) is a statutory body established on 1 January "
                "2003 under the AITI Order, 2001, and is an important regulator and development agency for Brunei's infocommunications and postal sectors. "
                "Its official website lists key responsibilities including regulatory functions for telecommunications and postal systems and services, "
                "development of the local ICT industry, protection of consumer interests through competitive pricing and service quality, management of "
                "the national radio-frequency spectrum, and administration and enforcement of the Personal Data Protection Order, 2025. AITI also acts "
                "as an adviser to government on policy and regulatory frameworks, a facilitator for collaboration and innovation, and a champion of "
                "development programmes supporting Digital Brunei. It therefore combines regulation, industry development, consumer protection, spectrum "
                "management and data-protection responsibilities, giving it an integrated platform role in Brunei's digital economy and communications "
                "market governance. "
                f"Source: {SOURCES['aiti']}"
            ),
        },
    ]

    for entry, update in zip(ENTRIES, expanded):
        entry.update(update)


def format_paragraph(paragraph):
    pf = paragraph.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    for run in paragraph.runs:
        set_run_font(run)


def add_label_paragraph(doc, label, body):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    label_run = paragraph.add_run(label)
    set_run_font(label_run, bold=True)
    body_run = paragraph.add_run(body)
    set_run_font(body_run)
    return paragraph


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal._element.rPr.rFonts.set(qn("w:cs"), FONT_NAME)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    title = doc.add_paragraph()
    title.paragraph_format.line_spacing = 1.5
    title.paragraph_format.space_after = Pt(12)
    title_run = title.add_run("2026APT学员单位调研报告")
    set_run_font(title_run, bold=True)

    note = doc.add_paragraph()
    note.paragraph_format.line_spacing = 1.5
    note.paragraph_format.space_after = Pt(12)
    note_run = note.add_run(
        "说明：本报告按名单中的13家单位整理，信息源均选用单位官网、政府门户、官方公报或官方通讯社等官方来源。"
    )
    set_run_font(note_run)

    apply_expanded_content()

    for idx, entry in enumerate(ENTRIES, start=1):
        if idx > 1:
            spacer = doc.add_paragraph()
            spacer.paragraph_format.line_spacing = 1.5
            spacer.paragraph_format.space_after = Pt(0)
        name_p = doc.add_paragraph()
        name_p.paragraph_format.line_spacing = 1.5
        name_p.paragraph_format.space_before = Pt(6)
        name_p.paragraph_format.space_after = Pt(6)
        name_run = name_p.add_run(f"{idx}. {entry['cn']}（{entry['en']}）")
        set_run_font(name_run, bold=True)

        add_label_paragraph(doc, "调研情况【中】：", entry["research_cn"])
        add_label_paragraph(doc, "调研情况【英】：", entry["research_en"])

    for paragraph in doc.paragraphs:
        format_paragraph(paragraph)

    doc.save(OUT_FILE)
    print(OUT_FILE.resolve())


if __name__ == "__main__":
    main()
